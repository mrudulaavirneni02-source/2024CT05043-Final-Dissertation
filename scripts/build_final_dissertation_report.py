from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUT = Path(__file__).resolve().parents[1] / "output"
OUT.mkdir(exist_ok=True)
DOCX = OUT / "2024CT05043_Final_Dissertation_Report_Updated.docx"
WORKFLOW_FIGURE = OUT / "logical_workflow.png"

TITLE = "INTELLIGENT KNOWLEDGE EXTRACTION FROM UNSTRUCTURED MULTIMODAL DATA USING RAG-BASED EVALUATION FRAMEWORKS"
STUDENT = "Avirneni Mrudula"
BITS_ID = "2024CT05043"
COURSE = "AMLCCZG628T - Dissertation"
PROGRAM = "M.Tech. Artificial Intelligence and Machine Learning"
ORGANIZATION = "Cognizant Technology Systems, Hyderabad"
SUPERVISOR = "Vidyasagar Parlapalli, Senior AI/ML Engineer, KLA"
ADDITIONAL_EXAMINER = "Prema Kumar Veerapaneni, Delivery Project Lead, Mphasis"
FACULTY_MENTOR = "To be completed before submission"
PROJECT_DURATION = "3 months"
DATE_OF_START = "23 May 2026"
DATE_OF_SUBMISSION = "23 August 2026"
DATE = "August 2026"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph, *, roman=False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    set_font(run, 10)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE \\* ROMAN" if roman else "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_para(doc, text="", *, align=None, bold=False, italic=False, size=12, before=0, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    if align is not None:
        p.alignment = align
    if keep:
        p.paragraph_format.keep_with_next = True
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size={1: 16, 2: 14, 3: 12}[level], bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r = p.add_run(text)
    set_font(r, 12)
    return p


def add_caption(doc, text, figure=False):
    p = add_para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11, before=3, after=8)
    p.paragraph_format.keep_with_next = False
    return p


def create_workflow_figure():
    width, height = 1500, 670
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("C:\\Windows\\Fonts\\timesbd.ttf", 34)
        text_font = ImageFont.truetype("C:\\Windows\\Fonts\\times.ttf", 25)
        small_font = ImageFont.truetype("C:\\Windows\\Fonts\\times.ttf", 21)
    except OSError:
        title_font = text_font = small_font = ImageFont.load_default()

    navy = "#17365D"
    blue = "#D9EAF7"
    green = "#E4F1E1"
    gold = "#FFF2CC"
    grey = "#F2F2F2"
    draw.text((width // 2, 22), "Intelligent Knowledge Extraction Framework", font=title_font, fill=navy, anchor="ma")

    boxes = [
        (60, 125, 300, 235, "Document upload\nand validation", blue),
        (365, 125, 605, 235, "Text extraction\nand chunking", blue),
        (670, 125, 910, 235, "ChromaDB vectors\n+ BM25 index", green),
        (975, 125, 1215, 235, "Vector / BM25 /\nhybrid retrieval", gold),
        (1280, 125, 1440, 235, "Grounding\ncheck", gold),
        (455, 415, 710, 530, "Local Llama 3.1\nthrough Ollama", grey),
        (850, 415, 1105, 530, "Cited answer,\nanalytics, evaluation", green),
    ]

    def rounded_box(x1, y1, x2, y2, text, fill):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=fill, outline=navy, width=3)
        lines = text.split("\n")
        y = (y1 + y2) / 2 - (len(lines) - 1) * 18
        for line in lines:
            draw.text(((x1 + x2) / 2, y), line, font=text_font, fill="black", anchor="mm")
            y += 36

    for box in boxes:
        rounded_box(*box)

    def arrow(start, end):
        draw.line((start, end), fill=navy, width=5)
        draw.polygon([(end[0], end[1]), (end[0] - 17, end[1] - 10), (end[0] - 17, end[1] + 10)], fill=navy)

    for x in (300, 605, 910, 1215):
        arrow((x + 4, 180), (x + 60, 180))
    arrow((1360, 240), (705, 415))
    arrow((710, 472), (850, 472))
    draw.text((750, 615), "All document processing and generation remain within the local deployment environment.", font=small_font, fill=navy, anchor="ma")
    image.save(WORKFLOW_FIGURE)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_widths(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, value in zip(header.cells, headers):
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1
        r = p.add_run(value)
        set_font(r, 10, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1
            r = p.add_run(str(value))
            set_font(r, 10)
    add_para(doc, "", after=2)
    return table


def add_cover(doc, title_page=False):
    for _ in range(5 if not title_page else 3):
        add_para(doc, "", after=0)
    add_para(doc, "A REPORT", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, after=8)
    add_para(doc, "ON", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, after=18)
    add_para(doc, TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=17, after=32)
    add_para(doc, "BY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, after=10)
    add_para(doc, STUDENT, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=15, after=4)
    add_para(doc, f"BITS ID: {BITS_ID}", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=4)
    add_para(doc, PROGRAM, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=30)
    if title_page:
        add_para(doc, "Prepared in partial fulfilment of the requirements of", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=3)
        add_para(doc, "WILP Dissertation Course", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=3)
        add_para(doc, COURSE, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, after=25)
    else:
        add_para(doc, "AT", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, after=8)
    add_para(doc, ORGANIZATION, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=42)
    add_para(doc, "BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, after=8)
    add_para(doc, DATE, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=0)
    doc.add_page_break()


def restart_page_numbering(section, start=1):
    sect_pr = section._sectPr
    page_numbers = sect_pr.find(qn("w:pgNumType"))
    if page_numbers is None:
        page_numbers = OxmlElement("w:pgNumType")
        sect_pr.append(page_numbers)
    page_numbers.set(qn("w:start"), str(start))


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(9)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(6)
    for level, size in ((1, 16), (2, 14), (3, 12)):
        style = styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    section.different_first_page_header_footer = True
    restart_page_numbering(section, start=0)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("Intelligent Knowledge Extraction from Unstructured Multimodal Data")
    set_font(run, 9, italic=True)
    add_page_number(section.footer.paragraphs[0], roman=True)

    add_cover(doc)
    add_cover(doc, title_page=True)

    add_heading(doc, "ACKNOWLEDGEMENTS", 1)
    add_para(doc, "I express my sincere gratitude to my project supervisor, Vidyasagar Parlapalli, Senior AI/ML Engineer, KLA, for his guidance, technical suggestions, and encouragement during this dissertation. His feedback helped shape the system architecture and the evaluation approach presented in this report.")
    add_para(doc, "I am thankful to the faculty members of the Work Integrated Learning Programmes Division, Birla Institute of Technology & Science, Pilani, for providing the academic framework and review process for this work. I also acknowledge Cognizant Technology Systems, Hyderabad, for the professional context that motivated the study of privacy-aware knowledge extraction from enterprise documents.")
    add_para(doc, "Finally, I thank my family and colleagues for their support throughout the implementation, testing, documentation, and refinement of the dissertation work.")
    doc.add_page_break()

    add_heading(doc, "ABSTRACT SHEET", 1)
    for label, value in [
        ("Organization and location", ORGANIZATION),
        ("Project duration", PROJECT_DURATION),
        ("Date of start", DATE_OF_START),
        ("Date of submission", DATE_OF_SUBMISSION),
        ("Course", COURSE),
        ("Title", TITLE.title()),
        ("Student", f"{STUDENT} ({BITS_ID})"),
        ("Supervisor", SUPERVISOR),
        ("Additional examiner", ADDITIONAL_EXAMINER),
        ("Faculty mentor", FACULTY_MENTOR),
        ("Project areas", "Artificial Intelligence; Natural Language Processing; Information Retrieval"),
        ("Key words", "Retrieval-Augmented Generation, hybrid retrieval, semantic search, BM25, ChromaDB, Ollama, evaluation"),
    ]:
        p = add_para(doc, before=0, after=1, size=10)
        r = p.add_run(f"{label}: ")
        set_font(r, size=10, bold=True)
        r = p.add_run(value)
        set_font(r, size=10)
    add_heading(doc, "Abstract", 2)
    add_para(doc, "Enterprise repositories contain reports, spreadsheets, manuals, and other documents whose value is difficult to access through keyword search alone. This dissertation develops an Intelligent Knowledge Extraction Framework that applies retrieval-augmented generation to locally indexed enterprise documents. The system validates uploaded files, extracts text from PDF, DOCX, TXT, CSV, XLSX, and Markdown sources, creates semantic chunks with metadata, and indexes the chunks in ChromaDB and a BM25 lexical index. A hybrid retrieval layer ranks candidate evidence before a local Llama 3.1 model, served through Ollama, prepares a cited answer. The implementation also records operational analytics and provides an evaluation interface for comparison of vector, lexical, and hybrid retrieval. A post-implementation functional test of nine representative questions confirmed correct handling of several framework and financial-document questions and demonstrated appropriate abstention for an unsupported question. The test also exposed cross-document evidence mixing and vector-retrieval false negatives. The report therefore treats the observed figures as exploratory and recommends a larger labelled benchmark before making accuracy claims. The resulting prototype provides a modular, privacy-aware foundation for evidence-based knowledge access in document-intensive environments.")
    add_para(doc, "Student signature: ____________________                         Supervisor signature: ____________________", before=18, after=0)
    doc.add_page_break()

    add_heading(doc, "TABLE OF CONTENTS", 1)
    toc_items = [
        "1. Introduction ........................................................................ 1",
        "2. Literature Review ................................................................. 3",
        "3. Research Methodology and System Design ........................... 4",
        "4. Implementation of the Proposed Framework ....................... 7",
        "5. Testing, Evaluation and Discussion .................................... 9",
        "6. Conclusions and Future Work ........................................... 12",
        "References ........................................................................... 13",
        "Glossary ............................................................................... 14",
        "Appendix A. API and Data-Processing Overview .................... 15",
        "Appendix B. Exploratory Test Set .......................................... 15",
        "Appendix C. Pre-Submission Checklist ................................ 16",
    ]
    for item in toc_items:
        add_para(doc, item, before=0, after=2)
    add_heading(doc, "LIST OF TABLES", 2)
    for item in [
        "Table 3.1. Major technical specifications ................................ 6",
        "Table 4.1. Implemented modules and responsibilities .............. 8",
        "Table 5.1. Exploratory test set ............................................... 9",
        "Table 5.2. Observed retrieval latency and displayed confidence ... 10",
        "Table 5.3. Interpretation of observed behaviours .................... 11",
    ]:
        add_para(doc, item, before=0, after=2)
    add_heading(doc, "LIST OF FIGURES", 2)
    add_para(doc, "Figure 3.1. Logical workflow of the proposed framework .......... 5", before=0, after=2)
    main_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    main_section.page_width = Inches(9)
    main_section.page_height = Inches(11)
    main_section.top_margin = Inches(1)
    main_section.bottom_margin = Inches(1)
    main_section.left_margin = Inches(1)
    main_section.right_margin = Inches(1)
    main_section.header_distance = Inches(0.5)
    main_section.footer_distance = Inches(0.5)
    main_section.header.is_linked_to_previous = False
    main_section.footer.is_linked_to_previous = False
    restart_page_numbering(main_section)
    main_header = main_section.header.paragraphs[0]
    main_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_header_run = main_header.add_run("Intelligent Knowledge Extraction from Unstructured Multimodal Data")
    set_font(main_header_run, 9, italic=True)
    add_page_number(main_section.footer.paragraphs[0])

    add_heading(doc, "1. INTRODUCTION", 1)
    add_heading(doc, "1.1 Background", 2)
    add_para(doc, "Organizations increasingly depend on unstructured information distributed across reports, policies, financial statements, spreadsheets, and technical documents. Although this material often contains valuable operational knowledge, its usefulness is reduced when users must manually inspect many files to find a precise answer. Keyword search can retrieve exact terms, but it is less reliable when a question and a relevant document use different vocabulary. Conversely, semantic retrieval can identify related concepts but may miss identifiers, abbreviations, and numerical details. These complementary strengths motivate the use of a hybrid retrieval approach.")
    add_para(doc, "Retrieval-Augmented Generation (RAG) combines information retrieval with a generative language model. Instead of asking a model to answer from its general training knowledge, a RAG pipeline first retrieves document passages and then asks the model to answer from that evidence [1]. For enterprise use, the approach can improve traceability when retrieved passages, source names, and page-level metadata are retained. It can also support local deployment when data sensitivity prevents the transfer of internal documents to an external service.")
    add_heading(doc, "1.2 Problem Statement", 2)
    add_para(doc, "The problem addressed in this work is the creation of a practical, locally deployable system that can ingest heterogeneous enterprise documents, retrieve relevant evidence through both semantic and lexical methods, and generate concise answers that remain bounded by the indexed material. The system must also reject unsuitable files, preserve useful metadata, expose retrieval choices for comparison, and provide an evaluation path that does not confuse retrieval confidence with verified answer accuracy.")
    add_heading(doc, "1.3 Objectives", 2)
    for text in [
        "Design a modular RAG architecture for enterprise document ingestion, retrieval, and answer generation.",
        "Implement validation, duplicate detection, text extraction, metadata preservation, and semantic chunking for supported formats.",
        "Compare dense-vector retrieval, BM25 lexical retrieval, and a hybrid ranking approach.",
        "Generate locally hosted, evidence-constrained responses and expose sources to the user.",
        "Provide an evaluation workflow that captures retrieval metrics, latency, and response-grounding observations.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "1.4 Scope and Limitations", 2)
    add_para(doc, "The implemented prototype focuses on text-bearing PDF, DOCX, TXT, CSV, XLSX, and Markdown documents. OCR fallback is available for pages without usable native text. The work does not claim production-scale load testing, multilingual evaluation, or a statistically sufficient accuracy benchmark. The local model can still produce an unsupported or mixed-source answer if retrieval admits weakly related passages; such cases are treated as findings of the evaluation rather than hidden limitations.")
    add_heading(doc, "1.5 Contributions", 2)
    add_para(doc, "The main contribution is an end-to-end and inspectable framework that joins document quality control, local semantic and lexical retrieval, response generation, and a user-facing evaluation interface. A later refinement of the prototype reduced chunk size, changed hybrid ranking to reciprocal-rank fusion, introduced subject-token coverage checks before generation, and added a library re-indexing workflow. These changes were motivated by observed retrieval failures rather than assumed to guarantee accuracy.")
    add_heading(doc, "1.6 Organization of the Report", 2)
    add_para(doc, "Chapter 2 reviews the concepts informing the study. Chapter 3 describes the research method and the system design. Chapter 4 documents the implemented modules. Chapter 5 presents the exploratory tests and their interpretation. Chapter 6 concludes the work and identifies the next steps required for a rigorous benchmark study.")

    add_heading(doc, "2. LITERATURE REVIEW", 1)
    add_heading(doc, "2.1 Retrieval-Augmented Generation", 2)
    add_para(doc, "RAG was introduced as a method for combining a parametric generator with retrieved non-parametric evidence for knowledge-intensive language tasks [1]. The architecture is especially relevant when an answer should be linked to an external corpus that can be updated independently of the language model. In an enterprise setting, this separation allows documents to be indexed and refreshed without retraining the generator.")
    add_heading(doc, "2.2 Dense Semantic Retrieval", 2)
    add_para(doc, "Dense retrieval represents text as vectors so that semantically similar passages can be identified by proximity in an embedding space. Sentence-BERT demonstrated an efficient approach for producing sentence embeddings suitable for similarity search [2]. The BGE family provides general-purpose embedding models designed for retrieval tasks [3]. In this dissertation, BAAI/bge-base-en-v1.5 is used to encode document chunks and user queries for vector retrieval through ChromaDB.")
    add_heading(doc, "2.3 Lexical Retrieval and BM25", 2)
    add_para(doc, "BM25 is a probabilistic lexical-ranking method that considers the occurrence of query terms, document length, and corpus statistics [4]. Its value in an enterprise corpus is practical: product names, reportable segments, exact financial labels, IDs, and technical acronyms are often best matched by explicit terms. However, a lexical score is query-relative, not an answer-confidence probability. This distinction is important when displaying system results to users.")
    add_heading(doc, "2.4 Hybrid Retrieval and Reranking", 2)
    add_para(doc, "Hybrid retrieval combines lexical and semantic candidates so that exact terms and conceptual similarity can reinforce each other. A naïve weighted sum is vulnerable when the underlying scores have different scales. The final prototype therefore uses reciprocal-rank fusion, which combines ranks rather than treating BM25 and vector scores as directly comparable. Candidate relevance is then subjected to a coverage-based grounding check before generation.")
    add_heading(doc, "2.5 Local Models and Evaluation", 2)
    add_para(doc, "Local model serving can support privacy and operational control where documents should remain within a managed environment. Llama 3.1 provides openly available instruction models suitable for local inference [5], while Ollama provides a local execution interface [6]. RAG evaluation should separately assess retrieval quality, whether an answer is supported by retrieved context, and whether it matches an accepted reference answer. Precision@K, Recall@K, MRR, and hit rate are retrieval measures; faithfulness and answer correctness require evidence and, in the latter case, a verified reference answer. Treating a model's confidence indicator as any of these metrics is methodologically unsound.")

    add_heading(doc, "3. RESEARCH METHODOLOGY AND SYSTEM DESIGN", 1)
    add_heading(doc, "3.1 Methodology", 2)
    add_para(doc, "The work followed a design-and-evaluation methodology. First, requirements were identified from the need for searchable enterprise documents, local operation, evidence visibility, and support for commonly used file formats. Second, the system was implemented as independently testable services. Third, representative documents and questions were used to exercise ingestion, retrieval, generation, abstention, and measurement behaviour. Finally, observations from the first test cycle were used to revise the chunking, fusion, grounding, and re-indexing mechanisms.")
    add_heading(doc, "3.2 Logical Workflow", 2)
    add_para(doc, "The logical workflow is: document upload -> validation -> text extraction -> normalization and semantic chunking -> embedding and dual indexing -> vector/BM25/hybrid retrieval -> grounding decision -> local answer generation -> citations, analytics, and evaluation. The flow intentionally separates information preparation from answer generation. This separation makes it possible to inspect stored chunks, compare retrieval modes, and block a response when the evidence is insufficient.")
    create_workflow_figure()
    doc.add_picture(str(WORKFLOW_FIGURE), width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "Figure 3.1. Logical workflow of the proposed Intelligent Knowledge Extraction Framework", figure=True)
    add_heading(doc, "3.3 System Components", 2)
    add_para(doc, "The backend is implemented using FastAPI and exposes services for health monitoring, ingestion, retrieval, chat, evaluation, analytics, and index inspection. The front end is built with React and Vite. ChromaDB persists vector embeddings, while an in-process BM25 index stores tokenized text for lexical search. The local generation service is accessed through Ollama. The architecture uses Pydantic schemas, explicit cancellation handling, and structured error responses to keep service boundaries clear.")
    add_heading(doc, "3.4 Data Validation and Extraction", 2)
    add_para(doc, "Before indexing, the ingestion service checks file extension, file size, corruption indicators, duplicate content through SHA-256 hashing, and the availability of readable text. Native extraction is used where possible, and OCR fallback is available for image-based PDF pages. Extracted text is accompanied by metadata such as filename, file type, page number, section, upload date, and chunk identifier. This metadata supports later citation, filtering, and collection inspection.")
    add_heading(doc, "3.5 Chunking and Indexing", 2)
    add_para(doc, "The final configuration uses evidence-sized semantic chunks with a target of 700 words and 100 words of overlap. This replaced the much larger initial configuration because broad chunks could join unrelated sections of a document and encourage mixed answers. Each chunk is embedded using BAAI/bge-base-en-v1.5 and stored in a cosine-similarity ChromaDB collection. The same cleaned text is tokenized for BM25. The persisted library can be rebuilt through a dedicated re-index operation whenever an indexing configuration changes.")
    add_heading(doc, "3.6 Retrieval, Grounding, and Generation", 2)
    add_para(doc, "Vector mode searches the ChromaDB collection with the encoded question. BM25 mode returns lexical candidates. Hybrid mode obtains candidates from both methods and ranks them by reciprocal-rank fusion with a constant of 60. The final grounding decision combines a confidence estimate with coverage of subject-bearing query terms in the retrieved evidence. Generic research words such as 'study', 'students', and 'results' are excluded from coverage computation so that weak topical overlap is less likely to authorize generation. The answer prompt limits the model to supplied excerpts, requires excerpt citations, and directs it to return 'Not found in the indexed documents' when direct support is absent.")
    add_heading(doc, "3.7 Major Technical Specifications", 2)
    add_table(doc, ["Parameter", "Final implementation"], [
        ("Programming environment", "Python 3.12 backend; TypeScript/React front end"),
        ("Backend API", "FastAPI with Pydantic request and response schemas"),
        ("Supported input", "PDF, DOCX, TXT, CSV, XLSX, Markdown"),
        ("Embedding model", "BAAI/bge-base-en-v1.5"),
        ("Vector store", "ChromaDB with cosine similarity"),
        ("Lexical retrieval", "BM25"),
        ("Hybrid ranking", "Reciprocal-rank fusion of vector and BM25 candidates"),
        ("Chunk configuration", "700-word target; 100-word overlap; semantic boundaries"),
        ("Generator", "Llama 3.1 (8B) served locally through Ollama"),
        ("Default response control", "Temperature 0.2; bounded context and output length"),
        ("Evaluation support", "Precision@K, Recall@K, MRR, hit rate, latency, RAGAS-compatible metrics"),
    ], [3150, 6930])
    add_caption(doc, "Table 3.1. Major technical specifications of the final prototype")

    add_heading(doc, "4. IMPLEMENTATION OF THE PROPOSED FRAMEWORK", 1)
    add_heading(doc, "4.1 Document Ingestion and Quality Controls", 2)
    add_para(doc, "The ingestion module accepts a file through the web interface, stores it temporarily, and runs validation before any embedding is created. Duplicate documents are detected using a SHA-256 content hash recorded in a local registry. Unsupported, oversized, corrupt, password-protected, and empty files are rejected with clear status messages. This approach protects the index from repeated or unusable material and makes the resulting collection easier to audit.")
    add_heading(doc, "4.2 Retrieval Service", 2)
    add_para(doc, "The retrieval service maintains the ChromaDB client, the embedding model, and the BM25 corpus. It supports vector, BM25, and hybrid modes through a common search interface. Source filtering is available when a user knows the document that should contain the answer. The service also exposes document counts and a read-only collection inspection view, allowing an evaluator to check source names, metadata, chunk text, and token counts.")
    add_heading(doc, "4.3 Response Service", 2)
    add_para(doc, "The response service prepares a compact, question-aware excerpt from each selected chunk and sends the resulting context to the local model. Each excerpt is labelled with a source number, filename, page number, section, and retrieval score. The prompt requires short evidence-only answers with bracketed excerpt references. A generation request is skipped when the grounding policy determines that the retrieved evidence does not sufficiently cover the question subject.")
    add_heading(doc, "4.4 Evaluation and Analytics", 2)
    add_para(doc, "The evaluation service can run a set of questions against one retrieval mode and record retrieval latency, ranking statistics, and optional RAGAS-based scores. The final interface makes a critical distinction: a verified reference answer and relevant source filename must be supplied before retrieval precision, recall, MRR, and hit rate are interpreted as ground-truth measures. Without those labels, the implementation reports zero rather than silently deriving a perfect score from its own similarity threshold. Analytics records document counts, duplicate and rejection events, retrieval time, generation time, and displayed confidence.")
    add_heading(doc, "4.5 Implemented Modules", 2)
    add_table(doc, ["Module", "Responsibility", "Outcome"], [
        ("Validation", "Checks type, integrity, duplicates, readable content, and file constraints.", "Protects index quality."),
        ("Extraction", "Obtains machine-readable text and metadata from supported formats.", "Creates consistent inputs."),
        ("Chunking", "Segments text around semantic units with controlled overlap.", "Produces evidence-sized records."),
        ("Indexing", "Creates embeddings and a BM25 corpus; persists ChromaDB data.", "Supports semantic and lexical search."),
        ("Retrieval", "Runs vector, BM25, or reciprocal-rank-fused hybrid search.", "Returns inspectable candidate evidence."),
        ("Grounding", "Tests confidence and subject-term coverage before generation.", "Reduces unsupported answers."),
        ("Generation", "Uses Ollama-hosted Llama 3.1 with source-labelled excerpts.", "Produces concise cited responses."),
        ("Evaluation", "Records labelled retrieval metrics and optional answer-level scores.", "Enables comparative analysis."),
    ], [2050, 5150, 2880])
    add_caption(doc, "Table 4.1. Implemented modules and their responsibilities")

    add_heading(doc, "5. TESTING, EVALUATION AND DISCUSSION", 1)
    add_heading(doc, "5.1 Test Purpose", 2)
    add_para(doc, "Testing was performed to verify end-to-end operation after the library was re-indexed into 44 evidence-sized chunks. The test suite contains nine representative questions: framework-design questions, technical-configuration questions, factual questions from an uploaded Microsoft annual report, a question about an uploaded language-learning study, and one intentionally unsupported question about Google's 2025 revenue. Each question was run in hybrid, vector, and BM25 modes. The recorded values are retained as a functional comparison, not as a final accuracy claim, because a complete gold source and reference answer were not supplied for every question in this exploratory run.")
    add_heading(doc, "5.2 Exploratory Test Set", 2)
    add_table(doc, ["ID", "Question category", "Expected behaviour"], [
        ("Q1", "Framework design considerations", "Retrieve design and architecture description."),
        ("Q2", "Semantic and lexical integration", "Explain the hybrid retrieval mechanism."),
        ("Q3", "Unsupported financial question", "Abstain because Google revenue is absent."),
        ("Q4", "Language-learning study method", "Identify the study's data-collection method."),
        ("Q5", "Evaluation metrics", "List proposed retrieval measures."),
        ("Q6", "Microsoft financial fact", "Return the FY2025 revenue figure."),
        ("Q7", "Microsoft segments", "Return segments and operating-income measure."),
        ("Q8-Q9", "Technology stack", "Return ingestion, embedding, database, BM25, and LLM details."),
    ], [970, 3230, 5880])
    add_caption(doc, "Table 5.1. Exploratory test set used for functional comparison")
    add_heading(doc, "5.3 Observed Retrieval Performance", 2)
    add_para(doc, "The observed latencies show a clear operational difference between retrieval modes. BM25 was the fastest method for this small local collection because it did not need query embedding or vector-database search. Hybrid and vector retrieval added embedding and vector-search overhead. The hybrid average was disproportionately increased by one 1,204.22 ms observation for the Microsoft-segment question; the median is included to show the central behaviour more fairly.")
    add_table(doc, ["Mode", "Mean displayed confidence", "Mean latency (ms)", "Median latency (ms)", "Interpretation"], [
        ("Hybrid", "67.8%", "235.2", "129.6", "More robust candidate coverage, but one large latency outlier."),
        ("Vector", "81.7%", "117.1", "102.5", "Semantic matches were often useful, but two answerable questions abstained."),
        ("BM25", "90.2%", "7.2", "5.3", "Fastest for this corpus and strong on explicit terms and financial labels."),
    ], [1200, 1720, 1620, 1620, 3920])
    add_caption(doc, "Table 5.2. Observed comparison across nine questions. Displayed confidence is not validated answer accuracy.")
    add_heading(doc, "5.4 Findings", 2)
    add_para(doc, "The framework correctly returned a supported abstention for the Google-revenue question in all three modes. This is an important behaviour because a knowledge-extraction system should distinguish unavailable information from a plausible answer generated from model priors. The framework and technology-stack questions generally produced concise answers with source markers. The Microsoft revenue and segment questions also demonstrated the usefulness of lexical retrieval for exact factual material.")
    add_para(doc, "The test also identified limitations. For the language-learning data-collection question, hybrid and vector responses combined the semi-structured interview evidence with an unrelated short-form-video survey. This is a cross-document contamination error: the answer contains a relevant fact, but it also includes unsupported material. In addition, vector retrieval returned 'Not found in the indexed documents' for two questions whose information was present in the corpus. The displayed grounding status therefore must not be read as proof that each sentence is correct. It only indicates that the retrieval and coverage policy allowed a generation attempt.")
    add_table(doc, ["Observed behaviour", "Likely cause", "Implemented or recommended response"], [
        ("Mixed evidence in Q4", "Related generic terms allowed candidates from different studies.", "Use smaller chunks, source-aware filtering, and a stronger reranker; manually review answer-level faithfulness."),
        ("Vector false negatives in Q2 and Q7", "Embedding retrieval did not rank the required section highly enough.", "Tune candidate depth; evaluate model choice and add a reranker."),
        ("High hybrid latency outlier", "Embedding, storage state, or local contention can affect one request.", "Log component timings and report median as well as mean latency."),
        ("High confidence without proof", "Confidence is a retrieval heuristic, not a labelled quality metric.", "Require gold sources and reference answers for final metrics."),
    ], [2380, 3390, 4310])
    add_caption(doc, "Table 5.3. Interpretation of observed behaviours and responses")
    add_heading(doc, "5.5 Threats to Validity", 2)
    add_para(doc, "The exploratory test set is small and contains documents from different domains, including a dissertation abstract, research papers, and a financial report. It is useful for detecting functional failures but is insufficient for estimating general retrieval accuracy. Some questions can be answered by lexical overlap alone, which favours BM25 in this particular collection. RAGAS-compatible metrics are available in the software, but local structured-output parsing can fail or time out; these scores should be reported only when completed successfully and when the associated reference answers are verified. A final benchmark should include at least 20 to 30 labelled questions, relevant source files, accepted answers, answerable and unanswerable cases, and manual review of every error.")
    add_heading(doc, "5.6 Discussion", 2)
    add_para(doc, "The results support the feasibility of a locally deployed knowledge-extraction workflow, but they do not establish that hybrid retrieval is universally superior. In this corpus, BM25 achieved the lowest observed latency and provided strong factual responses. Hybrid retrieval remains valuable when semantic paraphrase or incomplete keyword overlap matters, provided its candidate selection and grounding checks are carefully calibrated. The practical conclusion is that the system should expose retrieval modes, preserve evidence, and evaluate them on a labelled dataset instead of selecting a method solely on displayed confidence.")

    add_heading(doc, "6. CONCLUSIONS AND FUTURE WORK", 1)
    add_heading(doc, "6.1 Conclusions", 2)
    add_para(doc, "This dissertation designed and implemented an Intelligent Knowledge Extraction Framework for locally processed enterprise documents. The completed prototype accepts multiple document formats, validates input quality, extracts and chunks content, creates semantic and lexical indexes, retrieves evidence using three modes, and generates source-linked answers through a local Llama 3.1 model. The system also provides a user interface for ingestion, questioning, index inspection, analytics, and evaluation.")
    add_para(doc, "The implementation work demonstrates that local RAG can be assembled into a transparent and modular workflow. The post-implementation refinement is also a substantive result: broad chunks and naïvely normalized retrieval scores were replaced with evidence-sized chunks, cosine vector indexing, reciprocal-rank fusion, subject-token coverage checks, response limits, and a reproducible re-index workflow. These revisions reduced known sources of misleading confidence, although they do not eliminate the need for answer-level verification.")
    add_para(doc, "The nine-question functional comparison establishes that the system can retrieve and cite several supported answers, abstain on an unsupported request, and expose meaningful differences among retrieval modes. It also demonstrates why final claims must be measured carefully: the language-study response mixed evidence and vector retrieval missed two answerable questions. The final contribution is therefore both an implemented framework and an evaluation-aware account of its current strengths and limitations.")
    add_heading(doc, "6.2 Recommendations for Future Work", 2)
    for text in [
        "Construct and maintain a labelled benchmark with 20 to 30 or more questions, reference answers, and relevant source files.",
        "Add a cross-encoder reranker or learned relevance model after initial retrieval to reduce cross-document contamination.",
        "Evaluate additional embedding models and tune candidate depth, chunk size, overlap, and fusion parameters with held-out questions.",
        "Add answer-level citation validation so that each factual sentence is checked against its cited excerpt.",
        "Record separate embedding, vector-search, BM25, reranking, and generation timings to diagnose latency outliers.",
        "Extend extraction to richer multimodal evidence such as document images, tables, and captions, with explicit provenance.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "REFERENCES", 1)
    references = [
        "[1] P. Lewis, E. Perez, A. Piktus, et al., 'Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks,' Advances in Neural Information Processing Systems, vol. 33, pp. 9459-9474, 2020.",
        "[2] N. Reimers and I. Gurevych, 'Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks,' Proceedings of EMNLP-IJCNLP, 2019.",
        "[3] Beijing Academy of Artificial Intelligence, 'BGE: BAAI General Embedding Model,' 2024. Available: https://github.com/FlagOpen/FlagEmbedding.",
        "[4] S. Robertson and H. Zaragoza, 'The Probabilistic Relevance Framework: BM25 and Beyond,' Foundations and Trends in Information Retrieval, vol. 3, no. 4, pp. 333-389, 2009.",
        "[5] Meta AI, 'Llama 3.1: Open Foundation and Instruction Models,' 2024. Available: https://www.llama.com/.",
        "[6] Ollama, 'Ollama Documentation.' Available: https://ollama.com/.",
        "[7] Chroma, 'Chroma Documentation.' Available: https://docs.trychroma.com/.",
        "[8] FastAPI, 'FastAPI Documentation.' Available: https://fastapi.tiangolo.com/.",
        "[9] React, 'React Documentation.' Available: https://react.dev/.",
        "[10] Ragas, 'Ragas Documentation: Evaluation Framework for Retrieval-Augmented Generation.' Available: https://docs.ragas.io/.",
    ]
    for ref in references:
        add_para(doc, ref, before=0, after=4, size=11)

    add_heading(doc, "GLOSSARY", 1)
    add_table(doc, ["Term", "Meaning"], [
        ("BM25", "A probabilistic lexical ranking method for information retrieval."),
        ("ChromaDB", "A vector database used to persist and search embeddings."),
        ("Embedding", "A numerical representation of text used for semantic similarity search."),
        ("Faithfulness", "The extent to which a generated answer is supported by the retrieved context."),
        ("Hit rate", "The proportion of queries for which a relevant result is retrieved."),
        ("LLM", "Large Language Model."),
        ("MRR", "Mean Reciprocal Rank; a measure of the position of the first relevant result."),
        ("RAG", "Retrieval-Augmented Generation; generation conditioned on retrieved evidence."),
        ("Reciprocal-rank fusion", "A rank-combination method that aggregates results from multiple retrievers."),
    ], [2590, 7490])

    doc.add_page_break()
    add_heading(doc, "APPENDIX A. API AND DATA-PROCESSING OVERVIEW", 1)
    add_para(doc, "The implementation exposes REST endpoints for health monitoring, seed ingestion, document upload, asynchronous ingestion-job status, retrieval search, chat, evaluation, analytics, analytics export, and ChromaDB inspection. Upload requests are placed in a temporary incoming area and processed by an ingestion job. The job reports progress through hashing, validation, extraction, semantic boundary detection, embedding, indexing, and BM25 rebuilding. Cancellation support is available for long-running operations.")
    add_para(doc, "During indexing, each chunk is assigned an identifier and metadata. The metadata includes filename, source, page number where available, section, document type, file type, author where available, upload date, chunk position, token count, and chunking strategy. This information is preserved in ChromaDB and is used to display source citations and to filter retrieval by an individual document.")
    add_heading(doc, "APPENDIX B. EXPLORATORY TEST QUESTIONS", 1)
    questions = [
        "What are the main design considerations of the proposed Intelligent Knowledge Extraction Framework?",
        "How does the proposed framework combine semantic retrieval and BM25 lexical retrieval?",
        "What was Google's total revenue in fiscal year 2025 according to the uploaded documents?",
        "How was data collected in the study on mobile devices for English language learning?",
        "What evaluation metrics are proposed for assessing the effectiveness of the RAG retrieval system?",
        "What was Microsoft's total revenue in fiscal year 2025?",
        "What are the three reportable business segments of Microsoft, and what primary profitability measure does management use to review them?",
        "What technologies are used for document ingestion, embedding generation, retrieval, and response generation in the proposed system?",
        "What embedding model, vector database, lexical retrieval method, and LLM are used in the current knowledge extraction system?",
    ]
    for i, question in enumerate(questions, 1):
        add_para(doc, f"{i}. {question}", before=0, after=3)
    add_heading(doc, "APPENDIX C. PRE-SUBMISSION CHECKLIST", 1)
    for item in [
        "Replace the faculty mentor placeholder on the abstract sheet.",
        "Obtain supervisor review and signatures where required.",
        "Run the institutional plagiarism check and address any similarity findings.",
        "Update the final date, table of contents, and page numbers in the chosen editor if required by the submission portal.",
        "Export the final reviewed report as searchable PDF and verify that the file is below 10 MB.",
    ]:
        add_bullet(doc, item)

    doc.core_properties.author = STUDENT
    doc.core_properties.title = TITLE.title()
    doc.core_properties.subject = "WILP Dissertation Final Report"
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build()
