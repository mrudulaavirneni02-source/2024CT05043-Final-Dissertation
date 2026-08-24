"""Build the user-facing RAG architecture and operating guide.

Run from the repository root with:
    .\\.venv\\Scripts\\python.exe docs\\build_architecture_guide.py
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "RAG_Architecture_and_Operating_Guide.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "6B7280"
CONTENT_WIDTH_DXA = 9360


def set_cell_shading(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_width(cell, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_properties = table._tbl.tblPr
    table_width = table_properties.find(qn("w:tblW"))
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    table_properties.append(indent)
    layout = table_properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths, strict=True):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.margin_top = 80
            cell.margin_bottom = 80


def add_page_number(paragraph) -> None:
    run = paragraph.add_run("Page ")
    run.font.color.rgb = RGBColor.from_string(MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_run(paragraph, text: str, *, bold: bool = False, color: str | None = None, italic: bool = False) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    paragraph.add_run(text)


def add_number(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.167
    paragraph.add_run(text)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    paragraph.add_run(text)


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    add_run(paragraph, f"{label}: ", bold=True, color=DARK_BLUE)
    paragraph.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(INK)
    for row_data in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_data, strict=True):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.add_run(value)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    subtitle = styles.add_style("Guide Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(14)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("2024CT05043 Knowledge Extraction RAG | Architecture Guide")
    header_run.font.name = "Calibri"
    header_run.font.size = Pt(8)
    header_run.font.color.rgb = RGBColor.from_string(MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(footer)


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Knowledge Extraction RAG")
    run.font.name = "Calibri"
    run.font.size = Pt(27)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)
    subtitle = doc.add_paragraph(style="Guide Subtitle")
    subtitle.add_run("Architecture, ingestion, retrieval, and operating guide")
    metadata = doc.add_paragraph()
    metadata.paragraph_format.space_after = Pt(18)
    add_run(metadata, "Scope: ", bold=True, color=DARK_BLUE)
    metadata.add_run("text-first document ingestion with image-only OCR fallback")
    add_callout(
        doc,
        "Core rule",
        "Native/selectable PDF text is extracted with PyMuPDF. Tesseract OCR is used only for a PDF page where no native text is available. It is not used to OCR every page.",
    )


def build_document() -> None:
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    add_heading(doc, "1. What the system does")
    add_body(
        doc,
        "This project is a local Retrieval-Augmented Generation (RAG) application. It accepts business documents, extracts their readable text, breaks that text into meaningful chunks, indexes the chunks for search, and uses the retrieved evidence to answer questions. It supports PDF, DOCX, TXT, Markdown, CSV, and Excel files."
    )
    add_body(
        doc,
        "Think of it as a research assistant with a filing system. Ingestion is the filing step: the system reads a document and creates searchable notes. Retrieval is the lookup step: it finds the best notes for a question. Generation is optional: Ollama writes an answer only after the system has found sufficiently relevant evidence."
    )

    add_heading(doc, "2. Architecture at a glance")
    add_table(
        doc,
        ["Layer", "Responsibility", "Main implementation"],
        [
            ["Web interface", "Accept uploads, show progress, search, chat, and analytics.", "frontend/src/App.tsx and frontend/src/api.ts"],
            ["API and jobs", "Validate requests, accept files, run ingestion in background jobs, expose search/chat endpoints.", "backend/app/api/routes.py and services/jobs.py"],
            ["Extraction", "Read native document text; apply OCR only to textless PDF pages.", "backend/app/services/extraction.py"],
            ["Chunking", "Group related text into larger semantic chunks with citation metadata.", "backend/app/services/text_processing.py"],
            ["Indexes", "Store semantic vectors in ChromaDB and keyword data in BM25.", "backend/app/services/retrieval.py"],
            ["Answering", "Use retrieved evidence and Ollama to produce a grounded response.", "backend/app/services/llm.py"],
        ],
        [1800, 3900, 3660],
    )
    add_body(doc, "The normal request path is: browser -> FastAPI -> extraction -> chunking -> ChromaDB/BM25. A question follows the reverse route: browser -> FastAPI -> ChromaDB/BM25 -> optional Ollama answer -> browser.")

    add_heading(doc, "3. Ingestion flow: from upload to searchable text")
    add_body(doc, "The following sequence runs when a user uploads a document through the web interface or the upload API.")
    for step in [
        "The browser sends the file to POST /api/documents/upload. The API checks the extension and maximum upload size, then stores the incoming file temporarily.",
        "A background ingestion job calculates a SHA-256 hash. The duplicate detector prevents the same file from being indexed twice by accident.",
        "The corruption checker validates the format. For example, it rejects an encrypted PDF, an empty file, or a broken DOCX before expensive extraction starts.",
        "The extraction service reads text appropriate to the file type. PDFs use PyMuPDF first; DOCX files read word/document.xml; spreadsheets and CSV files are converted to readable sheet/row lines; TXT and Markdown are read as text.",
        "The text processor removes excess whitespace and groups related paragraphs into semantic chunks. Each chunk gets source, page, section, filename, type, upload time, and token-count metadata.",
        "The retrieval service embeds each chunk for ChromaDB, adds the original text to the persistent collection, then rebuilds the BM25 keyword index. The uploaded file is retained only after indexing succeeds.",
    ]:
        add_number(doc, step)
    add_callout(doc, "Example", "If a 10-page PDF has selectable text on pages 1-8, a screenshot on page 9, and a decorative photo on page 10, pages 1-8 are read with PyMuPDF, page 9 is sent to Tesseract only if it has no embedded text, and page 10 is harmlessly skipped if OCR finds nothing. The document still indexes the text from the other pages.")

    add_heading(doc, "4. PDF extraction and OCR: exactly when each is used")
    add_table(
        doc,
        ["Situation", "What happens", "Why"],
        [
            ["Normal digital PDF", "PyMuPDF calls page.get_text('text') and uses the returned text.", "This is fast and preserves true selectable text without OCR errors."],
            ["Scanned/image-only page", "If PyMuPDF returns no readable text, the page is rendered to an image and Tesseract reads that image.", "Scans store letters as pixels, so native text extraction has nothing to return."],
            ["PDF page with photo/diagram only", "Tesseract may return no text. The page is skipped and extraction continues.", "Graphs and non-text images are outside the current text-only scope."],
            ["Damaged image or OCR error", "The exception is caught for that page, logged, and the remaining pages continue.", "One bad image must not fail the complete uploaded file."],
            ["PDF with both text and images", "Native text is kept. OCR is not run over the full page merely because it contains an image.", "This avoids duplicate text and unnecessary OCR work."],
        ],
        [2000, 3700, 3660],
    )
    add_body(doc, "The OCR setting is ENABLE_OCR_FALLBACK=true in .env. The configured executable is TESSERACT_CMD=C:\\Program Files\\Tesseract-OCR\\tesseract.exe and the language is OCR_LANGUAGE=eng. If Tesseract is unavailable, only the OCR fallback is disabled; native text extraction continues normally.")

    add_heading(doc, "5. Chunking and indexing")
    add_body(doc, "Chunking converts a document into retrieval-sized passages. The current default is 3,000 words per target chunk, with a semantic boundary threshold of 0.85. This deliberately creates fewer, larger chunks than the previous 1,600-word configuration, reducing unnecessary fragmentation for ordinary reports.")
    add_body(doc, "The semantic chunker compares adjacent paragraphs using the embedding model. It keeps related paragraphs together until the chunk reaches the target size. If the content is short or has too little paragraph structure, the fallback recursive chunker divides it by word windows. Metadata preserves the page number and detected section so results can cite their origin.")
    add_table(
        doc,
        ["Index", "How it finds material", "Useful example"],
        [
            ["ChromaDB vector index", "Compares the meaning of the question with the meaning of each chunk using embeddings.", "A question about 'speed of lookup' can find a passage that says 'retrieval latency' without using the exact same phrase."],
            ["BM25 keyword index", "Ranks chunks using exact and important word matches.", "A question containing a product code, section number, or unusual abbreviation can find the literal term."],
            ["Hybrid retrieval", "Combines vector and BM25 scores with weights from .env.", "The usual choice: it balances conceptual matches with exact facts."],
        ],
        [2050, 4050, 3260],
    )
    add_callout(doc, "Retrieval note", "More chunks are not automatically better. Very small chunks can split a thought across many records and increase ranking noise. Larger semantic chunks keep a complete idea together; the system still returns only the configured top results, TOP_K=4, for a question.")

    add_heading(doc, "6. Search, grounded answers, and analytics")
    add_body(doc, "POST /api/retrieval/search returns the best matching chunks without generating prose. POST /api/chat first performs the same retrieval, calculates confidence, and calls Ollama only when the confidence meets MIN_CONTEXT_SCORE. Returned citations include the source filename, page/section metadata, and score. This is the guardrail that reduces unsupported answers.")
    add_body(doc, "The analytics service records upload, duplicate, rejection, retrieval, chat, and evaluation events in logs/events.jsonl. The /api/analytics endpoint summarizes document and chunk counts, average latency, confidence, duplicate files, and rejected files. The /api/analytics/export.csv endpoint exports the raw events for analysis.")

    add_heading(doc, "7. Scripts and how to use them")
    add_table(
        doc,
        ["Script", "When to use it", "What it does"],
        [
            ["scripts/bootstrap.py", "First-time setup or dependency refresh.", "Creates the virtual environment if needed, installs Python requirements, and runs npm install for the frontend."],
            ["scripts/run_backend.py", "To run the API locally.", "Starts Uvicorn using the host/port from .env. In development it enables reload."],
            ["scripts/ingest_file.py <path>", "To ingest a file without using the browser.", "Runs the same validation, extraction, chunking, indexing, and duplicate detection service flow as an upload."],
            ["scripts/ingest_seed.py", "To load the supplied dissertation seed transcript.", "Indexes data/seed/2024CT05043_abstract_clean.txt into ChromaDB and BM25."],
            ["scripts/evaluate_sample.py", "To smoke-test retrieval quality.", "Runs three sample questions using hybrid retrieval and reports retrieved count and scores."],
            ["scripts/check_ollama.py", "Before using generated chat answers.", "Checks whether Ollama is reachable and whether the model named in .env is installed."],
        ],
        [2450, 2800, 4110],
    )
    add_heading(doc, "8. Running the application")
    add_body(doc, "The backend listens on http://127.0.0.1:8000 and the frontend development server listens on http://127.0.0.1:5173. Open the frontend address in a browser to upload documents, inspect chunk totals, run retrieval, chat, and view analytics.")
    for command in [
        ".\\.venv\\Scripts\\python.exe scripts\\run_backend.py",
        "cd frontend",
        "npm.cmd run dev",
    ]:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(command)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    add_callout(doc, "Operational caution", "Chunking settings apply to newly indexed documents. A document already in the persistent indexes keeps its existing chunks until it is deliberately reindexed. Restart the backend after changing .env so it reads the new configuration.")

    doc.core_properties.title = "Knowledge Extraction RAG Architecture and Operating Guide"
    doc.core_properties.subject = "Architecture and operational guide"
    doc.core_properties.author = "2024CT05043 Knowledge Extraction RAG"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
